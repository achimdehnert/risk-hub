# substances/services/substance_import.py
"""
Import-Service für Gefahrstoffe.

Importiert reale Gefahrstoffdaten aus JSON-Dateien in die Datenbank.
Unterstützt:
- Stammdaten (Name, CAS, EC, Beschreibung)
- GHS-Klassifikation (H-/P-Sätze, Piktogramme, Signalwort)
- TRGS 510 Lagerklassen
- Ex-Schutz-relevante Daten (Flammpunkt, Zündtemperatur, UEG/OEG)
- Upsert-Logik (update_or_create) für idempotenten Import
"""

from __future__ import annotations

import json
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from uuid import UUID

from django.db import transaction

from substances.models import (
    HazardStatementRef,
    Identifier,
    PictogramRef,
    PrecautionaryStatementRef,
    SdsRevision,
    Substance,
)

logger = logging.getLogger(__name__)

DATA_DIR = Path(__file__).resolve().parent.parent / "data"
DEFAULT_DATA_FILE = DATA_DIR / "common_substances.json"


@dataclass
class ImportStats:
    """Tracks import statistics."""

    created: int = 0
    updated: int = 0
    skipped: int = 0
    errors: list[str] = field(default_factory=list)

    @property
    def total(self) -> int:
        return self.created + self.updated + self.skipped

    def summary(self) -> str:
        lines = [
            f"Import: {self.total} Stoffe verarbeitet",
            f"  Neu: {self.created}",
            f"  Aktualisiert: {self.updated}",
            f"  Übersprungen: {self.skipped}",
        ]
        if self.errors:
            lines.append(f"  Fehler: {len(self.errors)}")
            for err in self.errors[:10]:
                lines.append(f"    - {err}")
        return "\n".join(lines)


class SubstanceImportService:
    """Importiert Gefahrstoffe aus strukturierten JSON-Daten."""

    def __init__(
        self,
        tenant_id: UUID,
        user_id: UUID | None = None,
    ) -> None:
        self.tenant_id = tenant_id
        self.user_id = user_id

    def import_from_file(
        self,
        path: Path | None = None,
        *,
        dry_run: bool = False,
    ) -> ImportStats:
        """Import substances from a JSON file.

        Args:
            path: Path to JSON file. Uses default if None.
            dry_run: If True, validate but don't write.

        Returns:
            ImportStats with counts and errors.
        """
        data_path = path or DEFAULT_DATA_FILE
        if not data_path.exists():
            raise FileNotFoundError(f"Datendatei nicht gefunden: {data_path}")

        with open(data_path, encoding="utf-8") as f:
            records: list[dict[str, Any]] = json.load(f)

        logger.info(
            "Importiere %d Stoffe aus %s (dry_run=%s)",
            len(records),
            data_path.name,
            dry_run,
        )
        return self._import_records(records, dry_run=dry_run)

    def import_from_records(
        self,
        records: list[dict[str, Any]],
        *,
        dry_run: bool = False,
    ) -> ImportStats:
        """Import substances from a list of dicts."""
        return self._import_records(records, dry_run=dry_run)

    def _import_records(
        self,
        records: list[dict[str, Any]],
        *,
        dry_run: bool = False,
    ) -> ImportStats:
        stats = ImportStats()

        for idx, record in enumerate(records, 1):
            name = record.get("name", f"<unbenannt #{idx}>")
            try:
                if dry_run:
                    self._validate_record(record)
                    stats.skipped += 1
                else:
                    created = self._upsert_substance(record)
                    if created:
                        stats.created += 1
                    else:
                        stats.updated += 1
            except Exception as exc:
                msg = f"{name}: {exc}"
                logger.warning("Import-Fehler: %s", msg)
                stats.errors.append(msg)

        logger.info(stats.summary())
        return stats

    @transaction.atomic
    def _upsert_substance(
        self,
        record: dict[str, Any],
    ) -> bool:
        """Create or update a single substance. Returns True if created."""
        name = record["name"]

        substance, created = Substance.objects.update_or_create(
            tenant_id=self.tenant_id,
            name=name,
            defaults={
                "trade_name": record.get("trade_name", ""),
                "description": record.get("description", ""),
                "status": Substance.Status.ACTIVE,
                "storage_class": record.get("storage_class", ""),
                "is_cmr": record.get("is_cmr", False),
                "flash_point_c": record.get("flash_point_c"),
                "ignition_temperature_c": record.get("ignition_temperature_c"),
                "lower_explosion_limit": record.get("lower_explosion_limit"),
                "upper_explosion_limit": record.get("upper_explosion_limit"),
                "temperature_class": record.get("temperature_class", ""),
                "explosion_group": record.get("explosion_group", ""),
                "vapor_density": record.get("vapor_density"),
                "created_by": self.user_id,
            },
        )

        self._upsert_identifiers(substance, record)
        self._upsert_sds(substance, record)

        action = "erstellt" if created else "aktualisiert"
        logger.debug("Stoff %s: %s", name, action)
        return created

    def _upsert_identifiers(
        self,
        substance: Substance,
        record: dict[str, Any],
    ) -> None:
        """Create/update CAS, EC, and other identifiers."""
        id_map = {
            "cas": Identifier.IdType.CAS,
            "ec": Identifier.IdType.EC,
        }
        for key, id_type in id_map.items():
            value = record.get(key)
            if not value:
                continue
            Identifier.objects.update_or_create(
                tenant_id=self.tenant_id,
                substance=substance,
                id_type=id_type,
                defaults={
                    "id_value": value,
                    "created_by": self.user_id,
                },
            )

    def _upsert_sds(
        self,
        substance: Substance,
        record: dict[str, Any],
    ) -> None:
        """Create an initial SDS revision with H/P statements."""
        h_codes = record.get("h_statements", [])
        p_codes = record.get("p_statements", [])
        pictogram_codes = record.get("pictograms", [])
        signal_word = record.get("signal_word", "none")

        if not h_codes and not p_codes:
            return

        from django.utils import timezone

        sds, _ = SdsRevision.objects.update_or_create(
            substance=substance,
            revision_number=1,
            defaults={
                "tenant_id": self.tenant_id,
                "created_by": self.user_id,
                "revision_date": timezone.now().date(),
                "status": SdsRevision.Status.APPROVED,
                "signal_word": signal_word,
                "approved_by": self.user_id,
                "approved_at": timezone.now(),
                "notes": "Auto-Import aus Referenzdaten",
            },
        )

        if h_codes:
            refs = HazardStatementRef.objects.filter(code__in=h_codes)
            sds.hazard_statements.set(refs)

        if p_codes:
            refs = PrecautionaryStatementRef.objects.filter(code__in=p_codes)
            sds.precautionary_statements.set(refs)

        if pictogram_codes:
            refs = PictogramRef.objects.filter(code__in=pictogram_codes)
            sds.pictograms.set(refs)

    @staticmethod
    def _validate_record(record: dict[str, Any]) -> None:
        """Validate a record without writing to DB."""
        if not record.get("name"):
            raise ValueError("Feld 'name' fehlt")
        if not record.get("cas"):
            raise ValueError("Feld 'cas' fehlt")
        signal = record.get("signal_word", "none")
        valid_signals = {c.value for c in SdsRevision.SignalWord}
        if signal not in valid_signals:
            raise ValueError(f"Ungültiges Signalwort: {signal}")


    def import_from_csv(
        self,
        file_obj,
        *,
        dry_run: bool = False,
    ) -> ImportStats:
        """Import substances from a CSV file.

        Expected columns: name, cas, ec, trade_name, description,
        signal_word, h_statements (semicolon-separated),
        p_statements (semicolon-separated), pictograms (semicolon-separated),
        storage_class, is_cmr, flash_point_c, ignition_temperature_c,
        lower_explosion_limit, upper_explosion_limit, temperature_class,
        explosion_group, vapor_density
        """
        import csv
        import io

        content = file_obj.read()
        if isinstance(content, bytes):
            content = content.decode("utf-8-sig")

        reader = csv.DictReader(io.StringIO(content), delimiter=";")
        records = []
        for row in reader:
            record = {
                "name": (row.get("name") or "").strip(),
                "cas": (row.get("cas") or "").strip(),
                "ec": (row.get("ec") or "").strip(),
                "trade_name": (row.get("trade_name") or "").strip(),
                "description": (row.get("description") or "").strip(),
                "signal_word": (row.get("signal_word") or "none").strip().lower(),
                "storage_class": (row.get("storage_class") or "").strip(),
                "is_cmr": (row.get("is_cmr") or "").strip().lower() in ("true", "1", "ja", "yes"),
                "temperature_class": (row.get("temperature_class") or "").strip(),
                "explosion_group": (row.get("explosion_group") or "").strip(),
            }
            # Numeric fields
            for num_field in ("flash_point_c", "ignition_temperature_c",
                              "lower_explosion_limit", "upper_explosion_limit",
                              "vapor_density"):
                val = (row.get(num_field) or "").strip()
                if val:
                    try:
                        record[num_field] = float(val.replace(",", "."))
                    except ValueError:
                        pass

            # Semicolon-separated list fields
            for list_field in ("h_statements", "p_statements", "pictograms"):
                raw = (row.get(list_field) or "").strip()
                if raw:
                    record[list_field] = [c.strip() for c in raw.split(";") if c.strip()]

            records.append(record)

        logger.info("CSV-Import: %d Zeilen gelesen", len(records))
        return self._import_records(records, dry_run=dry_run)

    def import_from_xlsx(
        self,
        file_obj,
        *,
        dry_run: bool = False,
    ) -> ImportStats:
        """Import substances from an Excel (.xlsx) file.

        Reads the first sheet. First row = headers (same as CSV columns).
        """
        import openpyxl

        wb = openpyxl.load_workbook(file_obj, read_only=True, data_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        wb.close()

        if not rows:
            raise ValueError("Excel-Datei ist leer")

        headers = [str(h or "").strip().lower() for h in rows[0]]
        records = []
        for row in rows[1:]:
            raw = dict(zip(headers, row))
            record = {
                "name": str(raw.get("name") or "").strip(),
                "cas": str(raw.get("cas") or "").strip(),
                "ec": str(raw.get("ec") or "").strip(),
                "trade_name": str(raw.get("trade_name") or raw.get("handelsname") or "").strip(),
                "description": str(raw.get("description") or raw.get("beschreibung") or "").strip(),
                "signal_word": str(raw.get("signal_word") or raw.get("signalwort") or "none").strip().lower(),
                "storage_class": str(raw.get("storage_class") or raw.get("lagerklasse") or "").strip(),
                "is_cmr": str(raw.get("is_cmr") or raw.get("cmr") or "").strip().lower() in ("true", "1", "ja", "yes"),
                "temperature_class": str(raw.get("temperature_class") or "").strip(),
                "explosion_group": str(raw.get("explosion_group") or "").strip(),
            }
            # Numeric fields
            for num_field in ("flash_point_c", "ignition_temperature_c",
                              "lower_explosion_limit", "upper_explosion_limit",
                              "vapor_density"):
                val = raw.get(num_field)
                if val is not None:
                    try:
                        record[num_field] = float(val)
                    except (ValueError, TypeError):
                        pass

            # Semicolon-separated list fields
            for list_field in ("h_statements", "p_statements", "pictograms"):
                val = str(raw.get(list_field) or "").strip()
                if val:
                    record[list_field] = [c.strip() for c in val.split(";") if c.strip()]

            if record["name"]:
                records.append(record)

        logger.info("Excel-Import: %d Zeilen gelesen", len(records))
        return self._import_records(records, dry_run=dry_run)

    def import_from_docx(
        self,
        file_obj,
        *,
        dry_run: bool = False,
    ) -> ImportStats:
        """Import substances from a Word (.docx) file.

        Reads the first table in the document. First row = headers.
        """
        import docx

        doc = docx.Document(file_obj)
        if not doc.tables:
            raise ValueError("Kein Tabelle im Word-Dokument gefunden")

        table = doc.tables[0]
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]

        if not rows:
            raise ValueError("Tabelle ist leer")

        headers = [h.lower().replace(" ", "_") for h in rows[0]]
        records = []
        for row in rows[1:]:
            raw = dict(zip(headers, row))
            record = {
                "name": (raw.get("name") or raw.get("stoffname") or "").strip(),
                "cas": (raw.get("cas") or raw.get("cas-nummer") or raw.get("cas_nummer") or "").strip(),
                "ec": (raw.get("ec") or "").strip(),
                "trade_name": (raw.get("trade_name") or raw.get("handelsname") or "").strip(),
                "description": (raw.get("description") or raw.get("beschreibung") or raw.get("verwendung") or "").strip(),
                "signal_word": (raw.get("signal_word") or raw.get("signalwort") or "none").strip().lower(),
                "storage_class": (raw.get("storage_class") or raw.get("lagerklasse") or "").strip(),
                "is_cmr": (raw.get("is_cmr") or raw.get("cmr") or "").strip().lower() in ("true", "1", "ja", "yes"),
            }
            # Semicolon-separated list fields
            for list_field in ("h_statements", "p_statements", "pictograms"):
                alt_names = {"h_statements": ["h-sätze", "h_sätze", "h-saetze"],
                             "p_statements": ["p-sätze", "p_sätze", "p-saetze"],
                             "pictograms": ["piktogramme", "ghs"]}
                val = raw.get(list_field, "")
                if not val:
                    for alt in alt_names.get(list_field, []):
                        val = raw.get(alt, "")
                        if val:
                            break
                if val:
                    record[list_field] = [c.strip() for c in val.replace(",", ";").split(";") if c.strip()]

            if record["name"]:
                records.append(record)

        logger.info("DOCX-Import: %d Zeilen gelesen", len(records))
        return self._import_records(records, dry_run=dry_run)

    def import_from_pdf(
        self,
        file_obj,
        *,
        dry_run: bool = False,
    ) -> ImportStats:
        """Import substances from a PDF file.

        Extracts tables using pdfplumber. First row = headers.
        Falls back to text extraction if no tables found.
        """
        import pdfplumber

        content = file_obj.read()
        if isinstance(content, str):
            raise ValueError("PDF muss als Binärdatei übergeben werden")

        import io
        pdf = pdfplumber.open(io.BytesIO(content))
        all_rows = []
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                all_rows.extend(table)
        pdf.close()

        if not all_rows:
            raise ValueError(
                "Keine Tabellen in PDF gefunden. "
                "Bitte Excel oder CSV verwenden."
            )

        # First row with content = headers
        headers = [str(h or "").strip().lower().replace(" ", "_") for h in all_rows[0]]
        records = []
        for row in all_rows[1:]:
            if not any(cell for cell in row):
                continue
            raw = dict(zip(headers, [str(c or "").strip() for c in row]))
            record = {
                "name": (raw.get("name") or raw.get("stoffname") or raw.get("bezeichnung") or "").strip(),
                "cas": (raw.get("cas") or raw.get("cas-nummer") or raw.get("cas_nummer") or raw.get("cas-nr.") or "").strip(),
                "trade_name": (raw.get("trade_name") or raw.get("handelsname") or "").strip(),
                "signal_word": (raw.get("signal_word") or raw.get("signalwort") or "none").strip().lower(),
                "storage_class": (raw.get("storage_class") or raw.get("lagerklasse") or "").strip(),
                "is_cmr": (raw.get("is_cmr") or raw.get("cmr") or "").strip().lower() in ("true", "1", "ja", "yes"),
            }
            for list_field in ("h_statements", "p_statements"):
                alt_names = {"h_statements": ["h-sätze", "h_sätze", "h-saetze", "h-sätze"],
                             "p_statements": ["p-sätze", "p_sätze", "p-saetze", "p-sätze"]}
                val = raw.get(list_field, "")
                if not val:
                    for alt in alt_names.get(list_field, []):
                        val = raw.get(alt, "")
                        if val:
                            break
                if val:
                    record[list_field] = [c.strip() for c in val.replace(",", ";").replace(" ", ";").split(";") if c.strip()]

            if record["name"]:
                records.append(record)

        logger.info("PDF-Import: %d Zeilen gelesen", len(records))
        return self._import_records(records, dry_run=dry_run)

    def import_from_upload(
        self,
        file_obj,
        filename: str,
        *,
        dry_run: bool = False,
    ) -> ImportStats:
        """Import from an uploaded file (auto-detect format by extension).

        Supported: .json, .csv, .xlsx, .xls, .docx, .pdf
        """
        ext = Path(filename).suffix.lower()

        if ext == ".json":
            content = file_obj.read()
            if isinstance(content, bytes):
                content = content.decode("utf-8")
            records = json.loads(content)
            return self.import_from_records(records, dry_run=dry_run)
        elif ext == ".csv":
            return self.import_from_csv(file_obj, dry_run=dry_run)
        elif ext in (".xlsx", ".xls"):
            return self.import_from_xlsx(file_obj, dry_run=dry_run)
        elif ext == ".docx":
            return self.import_from_docx(file_obj, dry_run=dry_run)
        elif ext == ".pdf":
            return self.import_from_pdf(file_obj, dry_run=dry_run)
        else:
            raise ValueError(
                f"Nicht unterstütztes Dateiformat: {ext}. "
                "Erlaubt: .csv, .json, .xlsx, .docx, .pdf"
            )
