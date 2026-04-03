# substances/services/ai_extraction.py
"""
KI-gestützte Extraktion von Gefahrstoffdaten aus beliebigen Dokumenten.

Nutzt aifw (ADR-095) für intelligente Wert-Erkennung aus PDFs, Word-Dokumenten
und Excel-Dateien — unabhängig von Spaltenformat oder Sprache.
"""

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any
from uuid import UUID

logger = logging.getLogger(__name__)

ACTION_SUBSTANCE_IMPORT = "substance_import"

SYSTEM_PROMPT = """Du bist ein Experte für Gefahrstoff-Management nach GefStoffV.
Deine Aufgabe: Extrahiere Gefahrstoff-Daten aus dem folgenden Dokumenttext.

Gib ein JSON-Array zurück mit einem Objekt pro Gefahrstoff. Jedes Objekt hat diese Felder:
- "name": Stoffname (PFLICHT)
- "cas": CAS-Nummer (z.B. "67-64-1")
- "ec": EC-Nummer (optional)
- "trade_name": Handelsname (optional)
- "description": Beschreibung / Verwendungszweck (optional)
- "signal_word": "danger" oder "warning" oder "none"
- "h_statements": Array von H-Satz-Codes (z.B. ["H225", "H319", "H336"])
- "p_statements": Array von P-Satz-Codes (z.B. ["P210", "P233"])
- "pictograms": Array von GHS-Codes (z.B. ["GHS02", "GHS07"])
- "storage_class": TRGS 510 Lagerklasse (z.B. "3" für entzündbare Flüssigkeiten)
- "is_cmr": true/false — krebserzeugend, keimzellmutagent oder reproduktionstoxisch
- "flash_point_c": Flammpunkt in °C (Zahl oder null)
- "ignition_temperature_c": Zündtemperatur in °C (Zahl oder null)
- "lower_explosion_limit": UEG in Vol.% (Zahl oder null)
- "upper_explosion_limit": OEG in Vol.% (Zahl oder null)
- "temperature_class": z.B. "T1", "T2" etc. (optional)
- "explosion_group": z.B. "IIA", "IIB", "IIC" (optional)
- "vapor_density": Dampfdichte relativ zu Luft (Zahl oder null)

Regeln:
1. Extrahiere NUR Stoffe, die im Text tatsächlich vorkommen.
2. Wenn ein Feld nicht im Text steht, setze null oder leeren String.
3. H-/P-Sätze: Normalisiere auf Standard-Codes (H225, nicht "H 225" oder "Leichtentzündlich").
4. signal_word: Erkenne auch "Gefahr"→"danger", "Achtung"→"warning".
5. is_cmr: Setze true wenn H340, H341, H350, H351, H360, H361 vorhanden.
6. Gib NUR valides JSON zurück, KEINE Erklärung drumherum.
7. Bei einem einzelnen Sicherheitsdatenblatt: extrahiere den einen Stoff.
8. Bei einer Liste/Tabelle: extrahiere alle Stoffe."""

USER_PROMPT_TEMPLATE = """Extrahiere alle Gefahrstoffdaten aus folgendem Dokument.
Dateiname: {filename}

--- DOKUMENTTEXT ---
{text}
--- ENDE ---

Antworte NUR mit dem JSON-Array. Keine Erklärung."""


def extract_text_from_file(file_obj, filename: str) -> str:
    """Extract raw text from uploaded file for AI processing."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf_text(file_obj)
    elif ext == ".docx":
        return _extract_docx_text(file_obj)
    elif ext in (".xlsx", ".xls"):
        return _extract_xlsx_text(file_obj)
    elif ext == ".csv":
        content = file_obj.read()
        if isinstance(content, bytes):
            content = content.decode("utf-8-sig")
        return content
    elif ext == ".json":
        content = file_obj.read()
        if isinstance(content, bytes):
            content = content.decode("utf-8")
        return content
    else:
        raise ValueError(f"Nicht unterstütztes Format: {ext}")


def _extract_pdf_text(file_obj) -> str:
    """Extract text from PDF using pdfplumber."""
    import io

    import pdfplumber

    content = file_obj.read()
    pdf = pdfplumber.open(io.BytesIO(content))
    pages_text = []
    for page in pdf.pages:
        text = page.extract_text() or ""
        # Also try to get table data
        tables = page.extract_tables()
        for table in tables:
            for row in table:
                if row:
                    text += "\n" + " | ".join(str(c or "") for c in row)
        pages_text.append(text)
    pdf.close()
    return "\n\n".join(pages_text)


def _extract_docx_text(file_obj) -> str:
    """Extract text from Word document including tables."""
    import docx

    doc = docx.Document(file_obj)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_xlsx_text(file_obj) -> str:
    """Extract text from Excel as pipe-delimited rows."""
    import openpyxl

    wb = openpyxl.load_workbook(file_obj, read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = [str(c or "") for c in row]
            if any(c.strip() for c in cells):
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def ai_extract_substances(
    file_obj,
    filename: str,
    tenant_id: UUID | str | None = None,
) -> list[dict[str, Any]]:
    """Extract substance data from a file using AI (aifw).

    Returns a list of dicts ready for SubstanceImportService.import_from_records().
    """
    from ai_analysis.llm_client import llm_complete_sync

    # 1. Extract raw text
    raw_text = extract_text_from_file(file_obj, filename)

    if not raw_text.strip():
        raise ValueError("Dokument enthält keinen lesbaren Text.")

    # Limit text to avoid token overflow (approx 15k tokens)
    max_chars = 40_000
    if len(raw_text) > max_chars:
        logger.warning(
            "Text truncated from %d to %d chars for AI extraction",
            len(raw_text),
            max_chars,
        )
        raw_text = raw_text[:max_chars]

    # 2. Call LLM
    prompt = USER_PROMPT_TEMPLATE.format(
        filename=filename,
        text=raw_text,
    )

    logger.info("AI extraction starting for %s (%d chars)", filename, len(raw_text))

    response = llm_complete_sync(
        prompt=prompt,
        system=SYSTEM_PROMPT,
        action_code=ACTION_SUBSTANCE_IMPORT,
        tenant_id=tenant_id,
    )

    # 3. Parse JSON from response
    records = _parse_llm_response(response)

    logger.info("AI extracted %d substances from %s", len(records), filename)
    return records


def _parse_llm_response(response: str) -> list[dict[str, Any]]:
    """Parse LLM response, handling markdown code blocks."""
    text = response.strip()

    # Strip markdown code fences if present
    if text.startswith("```"):
        lines = text.split("\n")
        # Remove first and last lines (```json and ```)
        lines = [line for line in lines if not line.strip().startswith("```")]
        text = "\n".join(lines).strip()

    try:
        data = json.loads(text)
    except json.JSONDecodeError as e:
        logger.error("AI response is not valid JSON: %s", e)
        logger.debug("Raw response: %s", text[:500])
        raise ValueError(
            "KI-Antwort konnte nicht als JSON geparst werden. "
            "Bitte versuchen Sie es erneut oder nutzen Sie CSV/Excel."
        ) from e

    if isinstance(data, dict):
        # Maybe the LLM wrapped it in {"substances": [...]}
        for key in ("substances", "data", "records", "stoffe", "gefahrstoffe"):
            if key in data and isinstance(data[key], list):
                return data[key]
        return [data]
    elif isinstance(data, list):
        return data
    else:
        raise ValueError("Unerwartetes JSON-Format von KI")
