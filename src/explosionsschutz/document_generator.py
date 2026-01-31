# src/explosionsschutz/document_generator.py
"""
Document Generator für Explosionsschutzdokumente

Generiert Word-Dokumente aus ExplosionConcept-Daten mit:
- Corporate Design
- Strukturierte Kapitel
- Inhaltsverzeichnis
- Zonenpläne und Maßnahmen
"""

import io
from datetime import datetime
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .models import ExplosionConcept, Equipment


class ExSchutzDocumentGenerator:
    """Generiert Explosionsschutzdokumente als Word-Datei."""

    def __init__(self, concept: ExplosionConcept):
        self.concept = concept
        self.document = None

    def create_document(self, template_path: Optional[str] = None) -> "Document":
        """
        Erstellt das Word-Dokument.

        Args:
            template_path: Pfad zur Word-Vorlage (optional)

        Returns:
            Document: python-docx Document Objekt
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx ist nicht installiert. "
                "Bitte 'pip install python-docx' ausführen."
            )

        if template_path:
            self.document = Document(template_path)
        else:
            self.document = Document()
            self._setup_default_styles()

        self._add_cover_page()
        self._add_table_of_contents()
        self._add_zone_section()
        self._add_measures_section()
        self._add_equipment_section()
        self._add_revision_history()

        return self.document

    def _setup_default_styles(self):
        """Richtet Standard-Styles ein."""
        styles = self.document.styles

        if "Heading 1" in styles:
            h1 = styles["Heading 1"]
            h1.font.size = Pt(16)
            h1.font.bold = True

        if "Heading 2" in styles:
            h2 = styles["Heading 2"]
            h2.font.size = Pt(14)
            h2.font.bold = True

    def _add_cover_page(self):
        """Fügt Deckblatt hinzu."""
        self.document.add_paragraph()
        self.document.add_paragraph()

        title = self.document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("EXPLOSIONSSCHUTZDOKUMENT")
        run.bold = True
        run.font.size = Pt(24)

        subtitle = self.document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("gemäß § 6 Abs. 9 GefStoffV")
        run.font.size = Pt(14)

        self.document.add_paragraph()
        self.document.add_paragraph()

        info = self.document.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"Konzept: {self.concept.title}").bold = True

        if self.concept.area:
            area = self.document.add_paragraph()
            area.alignment = WD_ALIGN_PARAGRAPH.CENTER
            area.add_run(f"Bereich: {self.concept.area.name}")

        self.document.add_paragraph()
        self.document.add_paragraph()
        self.document.add_paragraph()

        table = self.document.add_table(rows=4, cols=2)
        table.style = "Table Grid"

        meta_data = [
            ("Dokument-Nr.:", f"EX-{str(self.concept.id)[:8].upper()}"),
            ("Version:", str(self.concept.version)),
            ("Status:", self.concept.get_status_display()),
            ("Erstellt am:", datetime.now().strftime("%d.%m.%Y")),
        ]

        for i, (label, value) in enumerate(meta_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)

        self.document.add_page_break()

    def _add_table_of_contents(self):
        """Fügt Inhaltsverzeichnis hinzu."""
        self.document.add_heading("Inhaltsverzeichnis", level=1)

        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()

        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")

        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

        hint = self.document.add_paragraph()
        hint.add_run(
            "(Bitte Inhaltsverzeichnis in Word aktualisieren: "
            "Rechtsklick → Felder aktualisieren)"
        ).italic = True

        self.document.add_page_break()

    def _add_zone_section(self):
        """Fügt Zoneneinteilung hinzu."""
        self.document.add_heading("1. Zoneneinteilung", level=1)

        zones = self.concept.zones.all().order_by("zone_type")

        if zones.exists():
            self.document.add_paragraph(
                f"Für das Konzept '{self.concept.title}' wurden "
                f"{zones.count()} Zonen definiert:"
            )

            table = self.document.add_table(rows=1, cols=4)
            table.style = "Table Grid"

            headers = table.rows[0].cells
            headers[0].text = "Zone"
            headers[1].text = "Bezeichnung"
            headers[2].text = "Ausdehnung (h/v)"
            headers[3].text = "Begründung"

            for zone in zones:
                row = table.add_row().cells
                row[0].text = zone.get_zone_type_display()
                row[1].text = zone.name
                h = zone.extent_horizontal_m or "-"
                v = zone.extent_vertical_m or "-"
                row[2].text = f"{h}m / {v}m"
                row[3].text = zone.justification[:100] if zone.justification else "-"
        else:
            self.document.add_paragraph("Keine Zonen definiert.")

        self.document.add_page_break()

    def _add_measures_section(self):
        """Fügt Schutzmaßnahmen hinzu."""
        self.document.add_heading("2. Schutzmaßnahmen", level=1)

        measures = self.concept.measures.all().order_by("category", "title")

        categories = [
            ("primary", "2.1 Primäre Maßnahmen (Vermeidung)"),
            ("secondary", "2.2 Sekundäre Maßnahmen (Zündquellenvermeidung)"),
            ("tertiary", "2.3 Tertiäre Maßnahmen (Auswirkungsbegrenzung)"),
            ("organizational", "2.4 Organisatorische Maßnahmen"),
        ]

        for cat_key, cat_title in categories:
            cat_measures = measures.filter(category=cat_key)

            self.document.add_heading(cat_title, level=2)

            if cat_measures.exists():
                for m in cat_measures:
                    p = self.document.add_paragraph(style="List Bullet")
                    p.add_run(m.title).bold = True
                    if m.description:
                        p.add_run(f": {m.description[:200]}")
                    p.add_run(f" [Status: {m.get_status_display()}]")
            else:
                self.document.add_paragraph("Keine Maßnahmen definiert.")

        self.document.add_page_break()

    def _add_equipment_section(self):
        """Fügt Betriebsmittel-Übersicht hinzu."""
        self.document.add_heading("3. Betriebsmittel in Ex-Bereichen", level=1)

        equipment = Equipment.objects.filter(
            zone__concept=self.concept, status="active"
        ).select_related("equipment_type", "zone")

        if equipment.exists():
            table = self.document.add_table(rows=1, cols=4)
            table.style = "Table Grid"

            headers = table.rows[0].cells
            headers[0].text = "Bezeichnung"
            headers[1].text = "Typ/Kennzeichnung"
            headers[2].text = "Zone"
            headers[3].text = "Nächste Prüfung"

            for eq in equipment:
                row = table.add_row().cells
                row[0].text = eq.name
                row[1].text = eq.equipment_type.atex_marking if eq.equipment_type else "-"
                row[2].text = eq.zone.get_zone_type_display() if eq.zone else "-"
                row[3].text = (
                    eq.next_inspection_date.strftime("%d.%m.%Y")
                    if eq.next_inspection_date
                    else "-"
                )
        else:
            self.document.add_paragraph("Keine Betriebsmittel erfasst.")

        self.document.add_page_break()

    def _add_revision_history(self):
        """Fügt Revisionsverlauf hinzu."""
        self.document.add_heading("Revisionsverlauf", level=1)

        table = self.document.add_table(rows=2, cols=4)
        table.style = "Table Grid"

        headers = table.rows[0].cells
        headers[0].text = "Version"
        headers[1].text = "Datum"
        headers[2].text = "Autor"
        headers[3].text = "Änderungen"

        row = table.rows[1].cells
        row[0].text = str(self.concept.version)
        row[1].text = datetime.now().strftime("%d.%m.%Y")
        row[2].text = (
            self.concept.created_by.get_full_name()
            if self.concept.created_by
            else "N/A"
        )
        row[3].text = "Erstversion" if self.concept.version == 1 else "Aktualisierung"

    def get_html_preview(self) -> str:
        """Generiert HTML-Vorschau des Dokuments."""
        zones = self.concept.zones.all()
        measures = self.concept.measures.all()

        html = f"""
        <style>
            .ex-doc {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; }}
            .ex-doc h1 {{ color: #0d6efd; border-bottom: 2px solid #0d6efd; }}
            .ex-doc h2 {{ color: #333; margin-top: 30px; }}
            .ex-doc .cover {{ text-align: center; padding: 50px 0; border: 1px solid #ddd; }}
            .ex-doc table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
            .ex-doc td, .ex-doc th {{ border: 1px solid #ddd; padding: 8px; }}
            .ex-doc th {{ background: #f8f9fa; }}
        </style>
        <div class="ex-doc">
            <div class="cover">
                <h1>EXPLOSIONSSCHUTZDOKUMENT</h1>
                <p>gemäß § 6 Abs. 9 GefStoffV</p>
                <h2>{self.concept.title}</h2>
                <p>Version {self.concept.version} | Status: {self.concept.get_status_display()}</p>
            </div>

            <h2>1. Zoneneinteilung</h2>
            <table>
                <tr><th>Zone</th><th>Bezeichnung</th><th>Begründung</th></tr>
        """

        for zone in zones:
            html += f"""
                <tr>
                    <td>{zone.get_zone_type_display()}</td>
                    <td>{zone.name}</td>
                    <td>{zone.justification[:100] if zone.justification else '-'}</td>
                </tr>
            """

        html += """
            </table>

            <h2>2. Schutzmaßnahmen</h2>
            <table>
                <tr><th>Kategorie</th><th>Maßnahme</th><th>Status</th></tr>
        """

        for measure in measures:
            html += f"""
                <tr>
                    <td>{measure.get_category_display()}</td>
                    <td>{measure.title}</td>
                    <td>{measure.get_status_display()}</td>
                </tr>
            """

        html += """
            </table>
        </div>
        """

        return html

    def save_to_buffer(self) -> io.BytesIO:
        """Speichert Dokument in BytesIO Buffer für Download."""
        if not self.document:
            self.create_document()

        buffer = io.BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer

    def get_filename(self) -> str:
        """Generiert Dateinamen für Download."""
        title = self.concept.title
        title = "".join(c for c in title if c.isalnum() or c in (" ", "-", "_")).strip()
        title = title.replace(" ", "_")
        date = datetime.now().strftime("%Y%m%d")
        return f"Explosionsschutzdokument_{title}_{date}.docx"


def generate_exschutz_document(concept_id: str) -> io.BytesIO:
    """
    Convenience-Funktion zum Generieren eines Explosionsschutzdokuments.

    Args:
        concept_id: UUID des Konzepts

    Returns:
        BytesIO Buffer mit Word-Dokument
    """
    concept = ExplosionConcept.objects.get(id=concept_id)
    generator = ExSchutzDocumentGenerator(concept)
    generator.create_document()
    return generator.save_to_buffer()
